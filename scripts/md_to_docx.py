from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MD_FILES = [
    ROOT / "README.md",
    ROOT / "docs" / "LaunchPad_Cursor_Build_Prompt.md",
    ROOT / "docs" / "ARCHITECTURE.md",
    ROOT / "docs" / "AGENTS.md",
    ROOT / "docs" / "WORKFLOWS.md",
    ROOT / "docs" / "TECH_FLOW.md",
    ROOT / "docs" / "DATABASE.md",
    ROOT / "docs" / "SPONSORS_AND_APIFY.md",
    ROOT / "docs" / "TEST_CASES.md",
]


def add_line(doc: Document, line: str) -> None:
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        return
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        return
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        return
    if line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
        return
    if line.startswith("|"):
        doc.add_paragraph(line)
        return
    doc.add_paragraph(line)


def convert(md_path: Path) -> Path:
    content = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    in_code = False

    for raw in content:
        line = raw.rstrip("\n")
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            doc.add_paragraph(line, style="Intense Quote")
            continue
        add_line(doc, line)

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main() -> None:
    generated = []
    for path in MD_FILES:
        if path.exists():
            generated.append(convert(path))
    for g in generated:
        print(g)


if __name__ == "__main__":
    main()
